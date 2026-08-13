#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
lies_contrl_ahb.py — liest die tabellarische Darstellung des BDEW-AHB CONTRL und
erzeugt `_engine/daten/contrl-ahb.js`.

Hintergrund (Protokoll Abschnitt 75): Die Codelisten selbst stehen im MIG und
werden von `lies_contrl_fehlercodes.py` gelesen (Abschnitt 74). Der AHB liefert
darüber hinaus die Angaben, die der MIG NICHT enthält:

  1. **Zulässigkeit je Anwendungsfall.** Der AHB führt drei Anwendungsfälle
     nebeneinander — „Empfangsbestätigung", „Syntaxfehlermeldung in der
     Übertragungsdatei" und „Syntaxfehlermeldung in der Nachricht" — und markiert
     je Code mit „X", in welchem davon er verwendet werden darf. Beispiel:
     UCI DE0083 = 7 gilt nur für die Empfangsbestätigung, = 4 nur für die beiden
     Fehlermeldungen. Die DE0085-Codes des UCI gehören ausschließlich zur
     Syntaxfehlermeldung in der ÜBERTRAGUNGSDATEI, die von UCM/UCS/UCD
     ausschließlich zur Syntaxfehlermeldung in der NACHRICHT.
  2. **AHB-Status samt Bedingungen** je Segment/Segmentgruppe/Datenelement
     (z. B. SG2 „Muss [9]", SG2 UCD „Soll [6]", SG1 UCM DE0085 „S [2] ∨ [3]")
     mitsamt den Bedingungstexten ([9] = „Wenn SG1 UCM DE0013 nicht vorhanden.").
  3. Die Bedingungstexte selbst.

Aufbau der Quelltabelle (verschachtelte Tabellen in einer Rahmentabelle; die
Zellen der äußeren Tabelle sind über `cell.text` LEER, weil python-docx
verschachtelte Tabellen dort nicht mitliest — deshalb `zellinhalt()`):

    Spalte 0/1 : "| SG1 | UCM | 0085 |"   Segmentgruppe, Segment, Datenelement
    Spalte 2   : "| 12 | Ungültiger Wert |  |  |  | X  | 13 | Fehlt | …"
    Spalte 3   : "[2] Wenn Syntaxfehler in UNH vorhanden. …"

In Spalte 2 folgen je Eintrag auf Code und Bezeichnung vier Felder: ein
Statusfeld und die drei Anwendungsfall-Spalten. Zeilen ohne eigenen Code führen
nur die Bezeichnung und danach dieselben vier Felder.

Aufruf:  python3 werkzeuge/lies_contrl_ahb.py [--pfad <AHB_CONTRL.docx>]
Ohne --pfad wird der AHB im Arbeitsordner gesucht (EDIGEN_ARBEITSORDNER).
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path

try:
    import docx  # python-docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:  # pragma: no cover
    sys.exit("python-docx fehlt: pip install python-docx")

REPO = Path(__file__).resolve().parents[1]
ARBEITSORDNER = Path(os.environ.get("EDIGEN_ARBEITSORDNER", REPO.parent.parent))
ZIEL = REPO / "_engine" / "daten" / "contrl-ahb.js"

# "| SG1 | UCM | 0085 |" bzw. "|  | UCI | 0020 |" / "|  | UCI |  | 00002 |"
KENNUNG = re.compile(
    r"^\|\s*(SG\d)?\s*\|\s*([A-Z]{3})\s*\|\s*(\d{4})?\s*\|?\s*(\d{5})?\s*\|?\s*$"
)
# Die Segmentgruppe selbst steht als nackte Kennung in einer eigenen Zeile
# ("SG2" mit Status "Muss [9]"). Das ist der GRUPPENSTATUS — genau die Angabe,
# die den Formular-Metas der Servicenachrichten fehlt (Protokoll Abschnitt 74).
GRUPPENKENNUNG = re.compile(r"^(SG\d)$")
BEDINGUNG = re.compile(r"\[(\d+)\]\s*([^\[]+)")
# Statusfeld: "X", "M", "S", "Muss", "Soll", "Kann" — jeweils optional mit
# Bedingungsausdruck ("S [2] ∨ [3]"). BEWUSST verankert und auf diese Wörter
# begrenzt: Ein Muster wie "S.*" würde auch Bezeichnungen verschlucken, die mit
# S beginnen ("Syntax-Version oder -ebene nicht unterstützt") — genau daran
# fehlte beim ersten Lauf der Code 2 in der UCI-Codeliste.
STATUSFELD = re.compile(r"(?:X|M|S|Muss|Soll|Kann)(?:\s*\[.*)?$")


def zellinhalt(cell) -> str:
    """Absätze UND verschachtelte Tabellen einer Zelle, in Dokumentreihenfolge."""
    teile = []
    for kind in cell._tc.iterchildren():
        tag = kind.tag.split("}")[1]
        if tag == "p":
            teile.append(Paragraph(kind, cell).text)
        elif tag == "tbl":
            for zeile in Table(kind, cell).rows:
                teile.append(" | ".join(zellinhalt(z) for z in zeile.cells))
    return " ".join(t for t in teile if t.strip())


def saeubere(text: str) -> str:
    """Trennstriche und Layout-Umbrüche der Kopfzellen auflösen."""
    text = re.sub(r"\s*\n\s*", " ", text)
    text = re.sub(r"(\w)-(\w)", r"\1\2", text)   # "Empfangs-bestätigung"
    return re.sub(r"\s{2,}", " ", text).strip()


def finde_ahb(pfad_arg: str | None) -> Path:
    if pfad_arg:
        p = Path(pfad_arg)
        if not p.is_file():
            sys.exit(f"AHB nicht gefunden: {p}")
        return p
    treffer = sorted(ARBEITSORDNER.rglob("AHB_CONTRL*.docx"))
    if not treffer:
        sys.exit(
            "AHB_CONTRL*.docx nicht im Arbeitsordner gefunden.\n"
            f"  Gesucht unter: {ARBEITSORDNER}\n"
            "  Abhilfe: EDIGEN_ARBEITSORDNER setzen oder --pfad angeben."
        )
    return treffer[-1]


def finde_tabelle(dokument) -> tuple[object, list[str]]:
    """Die Strukturtabelle samt Anwendungsfall-Namen aus ihrer Kopfzeile."""
    for tabelle in dokument.tables:
        if not tabelle.rows:
            continue
        kopf = " ".join(zellinhalt(c) for c in tabelle.rows[0].cells)
        if "EDIFACT Struktur" not in kopf:
            continue
        namen: list[str] = []
        for teil in kopf.split("|"):
            t = saeubere(teil)
            # Die Anwendungsfall-Spalten stehen zwischen "Beschreibung" und
            # "Bedingung" und sind wegen verbundener Zellen mehrfach genannt.
            if t and t not in namen and (
                "Empfangsbestätigung" in t or "Syntaxfehlermeldung" in t
            ):
                namen.append(t)
        if namen:
            return tabelle, namen
    sys.exit("Strukturtabelle (EDIFACT Struktur) im AHB nicht gefunden.")


def felder(text: str) -> list[str]:
    """Die per "|" getrennten Felder einer Inhaltszelle.

    Verbundene Zellen liefert python-docx mehrfach hintereinander mit gleichem
    Text; solche direkten Wiederholungen werden zusammengefasst, sonst
    verschiebt sich die Spaltenzählung (betrifft u. a. die Statuszeile von
    SG1 UCM DE0085, deren Bezeichnung doppelt erscheint).

    NICHT zusammengefasst werden Status-/X-Felder: „| Muss | Muss | Muss" meint
    drei Anwendungsfälle mit demselben Status, nicht eine verbundene Zelle. Das
    Zusammenfassen aller Wiederholungen hätte hier die Segment- und
    Gruppenstatus verschluckt.
    """
    roh = [f.strip() for f in text.split("|")]
    out: list[str] = []
    for f in roh:
        if out and f and f == out[-1] and not STATUSFELD.fullmatch(f):
            continue
        out.append(f)
    return out


def lies_eintraege(inhalt: str, anzahl_af: int) -> tuple[list[dict], list[str]]:
    """Codeeinträge und die Angaben der Zeile selbst aus einer Inhaltszelle.

    Aufbau je Eintrag: `Code | Bezeichnung | <Füllfeld> | <AF1> | <AF2> | <AF3>`.
    Die Anwendungsfall-Spalten tragen entweder „X" (Code in diesem Anwendungsfall
    verwendbar) oder — bei Struktur-/Datenelementzeilen ohne eigenen Code — den
    AHB-Status samt Bedingung („Muss", „Soll [6]", „S [2] ∨ [3]", „M"). Beides
    steht in DERSELBEN Spalte; einen separaten Statusblock gibt es nicht.
    """
    f = felder(inhalt)
    eintraege: list[dict] = []
    i = 1 if (f and not f[0]) else 0   # führendes Leerfeld, Text beginnt mit "|"
    erster_code = len(f)
    while i < len(f):
        wert = f[i]
        if not wert:
            i += 1
            continue
        # Ein Code ist numerisch (12, 13 …) oder ein Kürzel (UNA, UNH, APERAK).
        ist_code = bool(re.fullmatch(r"\d{1,3}|[A-Z]{3,7}", wert))
        if ist_code and i + 1 < len(f) and f[i + 1] and not STATUSFELD.fullmatch(f[i + 1]):
            erster_code = min(erster_code, i)
            rest = f[i + 2: i + 2 + 1 + anzahl_af]
            af = [x.strip() for x in rest[1:1 + anzahl_af]]
            af += [""] * (anzahl_af - len(af))
            eintraege.append({"code": wert, "text": f[i + 1], "af": af})
            i += 2 + 1 + anzahl_af
            continue
        i += 1
    # Die Angaben der Zeile selbst (Segment-, Gruppen- oder DE-Status bzw. „X")
    # stehen in den LETZTEN anzahl_af Feldern vor dem ersten Codeeintrag. Diese
    # Regel trägt über alle vorkommenden Zeilenformen hinweg — mit Bezeichnung
    # („| Datenaustauschreferenz |  | X | X | X"), ohne Bezeichnung
    # („| Muss | Muss | Muss") und mit führenden Leerfeldern („|  |  | Muss [9]").
    kopf = [x.strip() for x in f[:erster_code]]
    zeilen_af = kopf[-anzahl_af:] if len(kopf) >= anzahl_af else kopf
    zeilen_af = [""] * (anzahl_af - len(zeilen_af)) + zeilen_af
    return eintraege, zeilen_af


def lies(ahb: Path) -> dict:
    dokument = docx.Document(str(ahb))
    tabelle, anwendungsfaelle = finde_tabelle(dokument)
    anzahl_af = len(anwendungsfaelle)

    struktur: list[dict] = []
    bedingungen: dict[str, str] = {}
    for zeile in tabelle.rows:
        zellen = [zellinhalt(c).strip() for c in zeile.cells]
        uniq = list(dict.fromkeys(z for z in zellen if z))
        if not uniq:
            continue
        kennung = next((KENNUNG.match(u) for u in uniq if KENNUNG.match(u)), None)
        # Bedingungstexte stehen in der letzten Spalte.
        for u in uniq:
            if u.lstrip().startswith("["):
                for nr, txt in BEDINGUNG.findall(u):
                    bedingungen.setdefault(nr, saeubere(txt).rstrip(". ") + ".")
        gruppe = next((GRUPPENKENNUNG.match(u) for u in uniq if GRUPPENKENNUNG.match(u)), None)
        if not kennung and gruppe:
            inhalt = next((u for u in uniq if "|" in u
                           and not u.lstrip().startswith("[")), "")
            _codes, gruppen_af = lies_eintraege(inhalt, anzahl_af)
            struktur.append({"sg": gruppe.group(1), "segment": "", "de": "",
                             "af": gruppen_af, "codes": []})
            continue
        if not kennung:
            continue
        sg, segment, de, _zaehler = kennung.groups()
        inhalt = next((u for u in uniq if u is not kennung.string and "|" in u
                       and not u.lstrip().startswith("[")), "")
        eintraege, zeilen_af = lies_eintraege(inhalt, anzahl_af)
        struktur.append({
            "sg": sg or "",
            "segment": segment,
            "de": de or "",
            "af": zeilen_af,
            "codes": eintraege,
        })
    return {
        "anwendungsfaelle": anwendungsfaelle,
        "struktur": struktur,
        "bedingungen": bedingungen,
        "hinweise": lies_hinweise(dokument),
        "quelle": ahb.name,
    }


# Der AHB trifft im Fließtext Aussagen zu einzelnen Fehlercodes, die über die
# Codeliste hinausgehen — etwa das ausdrückliche Verbot, Code 26 zu senden, wenn
# der Empfänger die Übertragungsdatei aus selbst verursachtem Grund erneut
# einspielt. Solche Absätze werden mitgelesen und dem Code zugeordnet, damit die
# Anwendung sie am Befund anzeigen kann.
FEHLERCODE_SATZ = re.compile(r"Fehlercode\s+(\d{1,3})\b")


def lies_hinweise(dokument) -> dict[str, list[str]]:
    hinweise: dict[str, list[str]] = {}
    for absatz in dokument.paragraphs:
        text = saeubere(absatz.text)
        if len(text) < 40:
            continue
        for code in set(FEHLERCODE_SATZ.findall(text)):
            liste = hinweise.setdefault(code, [])
            if text not in liste:
                liste.append(text)
    return hinweise


def schreibe(daten: dict) -> None:
    af = daten["anwendungsfaelle"]
    # Nur die Einträge mit Codes sind für das Programm interessant; die reinen
    # Struktur-/Statuszeilen werden separat als Statusliste geführt.
    codelisten: dict[str, dict[str, list]] = {}
    status: dict[str, dict] = {}
    for eintrag in daten["struktur"]:
        pfad = " ".join(x for x in (eintrag["sg"], eintrag["segment"]) if x)
        if not pfad:
            continue
        schluessel = pfad + (" " + eintrag["de"] if eintrag["de"] else "")
        if any(eintrag["af"]):
            status[schluessel] = eintrag["af"]
        for c in eintrag["codes"]:
            if not eintrag["de"]:
                continue
            ziel = codelisten.setdefault(pfad, {}).setdefault(eintrag["de"], [])
            # Dasselbe DE kann mehrfach vorkommen (UCI DE0007 einmal für den
            # Absender, einmal für den Empfänger). Zulässigkeiten dann vereinen,
            # statt den Code doppelt zu führen.
            vorhanden = next((z for z in ziel if z[0] == c["code"]), None)
            if vorhanden:
                vorhanden[2] = [a or b for a, b in zip(vorhanden[2], c["af"])]
                continue
            ziel.append([c["code"], c["text"], [bool(x) for x in c["af"]]])

    zeilen = [
        "// _engine/daten/contrl-ahb.js",
        "// " + "-" * 66,
        "// AHB-Angaben zur CONTRL: Zulässigkeit der Codes JE ANWENDUNGSFALL sowie",
        "// AHB-Status und Bedingungen je Segment/Datenelement.",
        "//",
        "// MASCHINELL ERZEUGT aus dem BDEW-AHB CONTRL durch",
        "// werkzeuge/lies_contrl_ahb.py — nicht von Hand pflegen.",
        f"// Quelle: {daten['quelle']}",
        "//",
        "// Ergänzt _engine/daten/uci-fehlercodes.js (Codelisten aus dem MIG): Der MIG",
        "// sagt, WELCHE Codes es je Segment gibt; der AHB sagt, in welchem",
        "// ANWENDUNGSFALL sie verwendet werden dürfen. Beispiel UCI DE0083 — Code 7",
        "// nur in der Empfangsbestätigung, Code 4 nur in den Fehlermeldungen.",
        "// " + "-" * 66,
        "var contrlAhb = {",
        '  "anwendungsfaelle": ' + json.dumps(af, ensure_ascii=False) + ",",
        '  // Codelisten je Segmentpfad und Datenelement:',
        '  //   [code, Bezeichnung, [zulässig je Anwendungsfall …]]',
        '  "codes": {',
    ]
    pfade = sorted(codelisten)
    for pi, pfad in enumerate(pfade):
        zeilen.append(f'    "{pfad}": {{')
        des = codelisten[pfad]
        for di, de in enumerate(sorted(des)):
            eintraege = ",\n".join(
                "        " + json.dumps(e, ensure_ascii=False) for e in des[de]
            )
            komma = "," if di + 1 < len(des) else ""
            zeilen.append(f'      "{de}": [')
            zeilen.append(eintraege)
            zeilen.append(f"      ]{komma}")
        zeilen.append("    }" + ("," if pi + 1 < len(pfade) else ""))
    zeilen.append("  },")
    zeilen.append('  // AHB-Status je Segment/DE UND Anwendungsfall (Muss/Soll/M/S/X,')
    zeilen.append('  // ggf. mit Bedingung; leer = in diesem Anwendungsfall nicht verwendet).')
    zeilen.append('  "status": {')
    schluessel = sorted(status)
    for si, s in enumerate(schluessel):
        komma = "," if si + 1 < len(schluessel) else ""
        zeilen.append(f'    "{s}": ' + json.dumps(status[s], ensure_ascii=False) + komma)
    zeilen.append("  },")
    zeilen.append('  // Aussagen des AHB-Fließtexts zu einzelnen Fehlercodes.')
    zeilen.append('  "hinweise": {')
    hin = daten.get("hinweise", {})
    hnum = sorted(hin, key=int)
    for hi, nr in enumerate(hnum):
        komma = "," if hi + 1 < len(hnum) else ""
        zeilen.append(f'    "{nr}": ' + json.dumps(hin[nr], ensure_ascii=False) + komma)
    zeilen.append("  },")
    zeilen.append('  "bedingungen": {')
    bed = daten["bedingungen"]
    nummern = sorted(bed, key=int)
    for bi, nr in enumerate(nummern):
        komma = "," if bi + 1 < len(nummern) else ""
        zeilen.append(f'    "{nr}": ' + json.dumps(bed[nr], ensure_ascii=False) + komma)
    zeilen.append("  }")
    zeilen.append("};")
    zeilen.append("")
    zeilen.append("// Ist `code` an dieser Stelle laut AHB vorgesehen? Rückgabe:")
    zeilen.append("//   { bekannt, anwendungsfaelle: [Namen], text }")
    zeilen.append("// bekannt === false heißt: Der AHB führt den Code an dieser Ebene")
    zeilen.append("// nicht — entweder falsche Ebene oder im Marktprozess nicht vorgesehen.")
    zeilen.append("function contrlAhbZulaessig(code, segmentPfad, de) {")
    zeilen.append('  var liste = ((contrlAhb.codes || {})[segmentPfad] || {})[de || "0085"] || [];')
    zeilen.append("  for (var i = 0; i < liste.length; i++) {")
    zeilen.append("    if (liste[i][0] !== String(code)) continue;")
    zeilen.append("    var namen = [];")
    zeilen.append("    (liste[i][2] || []).forEach(function (ja, k) {")
    zeilen.append("      if (ja && contrlAhb.anwendungsfaelle[k]) namen.push(contrlAhb.anwendungsfaelle[k]);")
    zeilen.append("    });")
    zeilen.append("    return { bekannt: true, anwendungsfaelle: namen, text: liste[i][1] };")
    zeilen.append("  }")
    zeilen.append('  return { bekannt: false, anwendungsfaelle: [], text: "" };')
    zeilen.append("}")
    zeilen.append("")
    zeilen.append('if (typeof module !== "undefined")')
    zeilen.append("  module.exports = { contrlAhb: contrlAhb, contrlAhbZulaessig: contrlAhbZulaessig };")
    zeilen.append("")
    ZIEL.write_text("\n".join(zeilen), encoding="utf-8")


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--pfad", help="Pfad zur AHB_CONTRL-DOCX (sonst Suche im Arbeitsordner)")
    args = ap.parse_args()

    ahb = finde_ahb(args.pfad)
    daten = lies(ahb)
    schreibe(daten)

    print(f"Quelle          : {ahb}")
    print(f"Anwendungsfälle : {len(daten['anwendungsfaelle'])}")
    for n in daten["anwendungsfaelle"]:
        print(f"                  - {n}")
    mit_codes = [e for e in daten["struktur"] if e["codes"]]
    print(f"Zeilen mit Codes: {len(mit_codes)}")
    for e in mit_codes:
        pfad = " ".join(x for x in (e["sg"], e["segment"], e["de"]) if x)
        codes = ", ".join(c["code"] for c in e["codes"])
        print(f"  {pfad:<16} {len(e['codes']):>2} Codes  {codes[:78]}")
    print(f"Bedingungen     : {len(daten['bedingungen'])}")
    print(f"Fließtexthinweise zu Codes: {sorted(daten['hinweise'], key=int)}")
    print(f"Ziel            : {ZIEL.relative_to(REPO)}")


if __name__ == "__main__":
    main()
