#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
tabs_ahb_reader.py — Leser für AHB-DOCX im klassischen Layout (tabulatorgetrennte Spalten).

Gegenstück zu `nested_ahb_reader.py`. Beide erzeugen dieselbe Datenstruktur, sodass
für alle Nachrichtentypen und beide Formatstände eine einheitliche Datenbasis entsteht.

Im klassischen Layout steckt die gesamte mittlere Spalte einer Tabellenzeile in einer
einzigen Zelle; die Spalten werden über Tabstop-Positionen der Absätze getrennt:

    Absatz 'Prüfidentifikator\\t55001\\t55002\\t55003' mit tabs=[a, b, c]
       -> Spaltenposition a gehört zu 55001, b zu 55002, c zu 55003

    Absatz '137\\tDokumenten-/\\tX\\tX\\tX' mit tabs=[name_pos, a, b, c]
       -> '137' steht links des ersten Tabstops und ist der Codewert,
          'Dokumenten-/' die Bezeichnung, danach je Prüf-ID der AHB-Ausdruck.

Die Tabstop-Positionen schwanken zwischen Kopf- und Inhaltszeilen um wenige tausend
EMU; die Zuordnung erfolgt deshalb über die geringste Abweichung.

Bearbeiter: Steffen Haense / Claude — 28.07.2026
"""

from __future__ import annotations

import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.table import _Cell

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

RE_SEGMENT_ID = re.compile(r"^\d{5}$")
RE_DATA_ELEMENT = re.compile(r"^\d{4}$")
RE_SEGMENT_GROUP = re.compile(r"^SG\d+$")
RE_SEGMENT_CODE = re.compile(r"^[A-Z]{3}$")
RE_PRUEFI = re.compile(r"^\d{5}$")

#: Zwei Tabstops gelten als dieselbe Spalte, wenn sie näher als dieser Wert liegen (EMU).
SPALTEN_TOLERANZ = 120_000


def _tabstops(paragraph) -> list[int]:
    return [ts.position for ts in paragraph.paragraph_format.tab_stops]


def _ist_codewert(text: str) -> bool:
    """Erkennt Codewerte (Value-Pool-Einträge) wie '137', 'ZW4', 'E01', 'GABi-RLMmT'."""
    try:
        from kohlrahbi.models.flat_ahb_reader import FlatAhbCsvReader

        return bool(FlatAhbCsvReader._is_value_pool_entry(candidate=text))  # noqa: SLF001
    except Exception:  # pragma: no cover — Rückfallregel ohne kohlrahbi
        return bool(text) and " " not in text and len(text) <= 20


@dataclass
class TableHead:
    pruefis: list[str]
    positionen: list[int]
    beschreibung: dict[str, str] = field(default_factory=dict)
    kommunikation: dict[str, str] = field(default_factory=dict)


def read_head(cell) -> TableHead | None:
    """Liest den Tabellenkopf einer klassischen AHB-Tabelle."""
    absaetze = cell.paragraphs
    if not absaetze:
        return None
    letzter = absaetze[-1]
    if not letzter.text.startswith("Prüfidentifikator"):
        return None
    teile = letzter.text.split("\t")[1:]
    positionen = _tabstops(letzter)
    pruefis = [t.strip() for t in teile if RE_PRUEFI.match(t.strip())]
    if not pruefis or len(positionen) < len(pruefis):
        return None
    head = TableHead(pruefis=pruefis, positionen=positionen[: len(pruefis)])

    abschnitt: str | None = None
    for p in absaetze:
        teile = p.text.split("\t")
        kopf = teile[0].strip()
        if kopf == "Prüfidentifikator":
            continue
        if kopf in ("Beschreibung", "Kommunikation von"):
            abschnitt = kopf
            ziel = head.beschreibung if kopf == "Beschreibung" else head.kommunikation
            for pruefi, text in zip(head.pruefis, teile[1:]):
                ziel[pruefi] = (ziel.get(pruefi, "") + " " + text.strip()).strip()
        elif abschnitt == "Beschreibung":
            # Fortsetzungszeile: über die Tabstop-Positionen zuordnen
            pos = _tabstops(p)
            werte = teile[1:] if teile and not teile[0].strip() else teile
            for tabstop, text in zip(pos, werte):
                if not text.strip():
                    continue
                idx = min(range(len(head.positionen)), key=lambda k: abs(head.positionen[k] - tabstop))
                if abs(head.positionen[idx] - tabstop) <= SPALTEN_TOLERANZ:
                    pruefi = head.pruefis[idx]
                    head.beschreibung[pruefi] = (head.beschreibung.get(pruefi, "") + " " + text.strip()).strip()
    for pruefi in head.pruefis:
        head.beschreibung[pruefi] = " ".join(head.beschreibung.get(pruefi, "").split())
        head.kommunikation[pruefi] = " ".join(head.kommunikation.get(pruefi, "").split())
    return head


def kalibriere(zellen_folge, head: TableHead) -> None:
    """Richtet die Spaltenpositionen an den Datenzeilen aus.

    Die Tabstops der Kopfzeile und die der Inhaltszeilen sind in einigen Dokumenten
    um einen konstanten Betrag gegeneinander verschoben (beobachtet: rund 268.000 EMU
    in AHB UTILMD Strom S2.2). Ohne Ausgleich landen sämtliche AHB-Ausdrücke in der
    Bezeichnungsspalte und die Prüf-ID bleibt ohne Zeilen. Deshalb werden die
    Spaltenpositionen aus der ersten hinreichend belegten Inhaltszeile übernommen:
    die n rechtesten Tabstops sind die n Prüf-ID-Spalten.
    """
    n = len(head.pruefis)
    for zelle in zellen_folge:
        for p in zelle.paragraphs:
            tabstops = sorted(_tabstops(p))
            if len(tabstops) >= n and p.text.count("\t") >= n:
                kandidat = tabstops[-n:]
                abstand = abs(kandidat[0] - head.positionen[0])
                if abstand > SPALTEN_TOLERANZ:
                    head.positionen = kandidat
                return


@dataclass
class StructInfo:
    segment_group_key: str | None = None
    segment_code: str | None = None
    data_element: str | None = None
    segment_id: str | None = None


def parse_struct(cell) -> StructInfo:
    info = StructInfo()
    texte = [t.strip() for t in cell.text.replace("\n", "\t").split("\t") if t.strip()]
    for t in texte:
        if RE_SEGMENT_GROUP.match(t):
            info.segment_group_key = t
        elif RE_SEGMENT_CODE.match(t) and info.segment_code is None:
            info.segment_code = t
        elif RE_SEGMENT_ID.match(t):
            info.segment_id = t
        elif RE_DATA_ELEMENT.match(t):
            info.data_element = t
        elif info.segment_code is not None and info.data_element is None:
            info.data_element = t
    return info


@dataclass
class BodyEntry:
    name: str = ""
    value_pool_entry: str | None = None
    expressions: list[str] = field(default_factory=list)


def parse_body(cell, head: TableHead) -> list[BodyEntry]:
    """Zerlegt die mittlere Spalte einer Datenzeile in Codewerte, Namen und Ausdrücke."""
    n = len(head.pruefis)
    if not cell.paragraphs or not cell.text.strip():
        return []

    entries: list[BodyEntry] = []

    def neue_zeile() -> BodyEntry:
        e = BodyEntry(expressions=[""] * n)
        entries.append(e)
        return e

    aktuell: BodyEntry | None = None

    for p in cell.paragraphs:
        text = p.text.replace("\xa0", "")
        if not text.strip():
            continue
        tabstops = _tabstops(p)
        teile = text.split("\t")

        # Positionen der Textstücke: das erste steht links des ersten Tabstops.
        if len(teile) == len(tabstops) + 1:
            positionen = [p.paragraph_format.left_indent or 0] + list(tabstops)
        elif len(teile) == len(tabstops):
            positionen = list(tabstops)
        else:
            positionen = ([p.paragraph_format.left_indent or 0] + list(tabstops))[: len(teile)]
            while len(positionen) < len(teile):
                positionen.append(positionen[-1] if positionen else 0)

        # Textstücke den Spalten zuordnen
        pid_teile: list[tuple[int, str]] = []
        vorspann: list[tuple[int, str]] = []
        for pos, stueck in zip(positionen, teile):
            idx = min(range(n), key=lambda k: abs(head.positionen[k] - pos))
            if abs(head.positionen[idx] - pos) <= SPALTEN_TOLERANZ:
                pid_teile.append((idx, stueck.strip()))
            else:
                vorspann.append((pos, stueck.strip()))

        # Steht links der Bezeichnungsspalte noch ein Textstück, ist das der Codewert.
        # Die Bezeichnungsspalte ist der Tabstop deutlich links der ersten Prüf-ID-Spalte.
        grenze = head.positionen[0] - SPALTEN_TOLERANZ
        bezeichnungs_tabstops = [t for t in tabstops if t < grenze]
        code: str | None = None
        name_teile: list[str] = []
        if bezeichnungs_tabstops and vorspann:
            erste_bezeichnung = min(bezeichnungs_tabstops)
            for pos, stueck in vorspann:
                if not stueck:
                    continue
                if pos < erste_bezeichnung - 1000 and code is None and not name_teile:
                    code = stueck
                else:
                    name_teile.append(stueck)
        else:
            name_teile = [t for _, t in vorspann if t]

        # Ein neuer Codewert beginnt nur dann eine neue Zeile, wenn der Absatz auch
        # AHB-Ausdrücke trägt. Absätze ohne Ausdruck setzen einen über den Zeilen-
        # umbruch getrennten Code bzw. Namen fort (z. B. "GABi-" + "RLMmT").
        hat_ausdruck = any(t for _, t in pid_teile)
        beginnt_neue_zeile = code is not None and aktuell is not None and hat_ausdruck

        if aktuell is None or beginnt_neue_zeile:
            aktuell = neue_zeile()
            aktuell.value_pool_entry = code
            aktuell.name = " ".join(name_teile)
        else:
            if code is not None:
                if aktuell.value_pool_entry is None:
                    aktuell.value_pool_entry = code
                else:
                    # Codewerte enthalten nie Leerzeichen: über den Zeilenumbruch
                    # getrennte Teile werden unmittelbar zusammengesetzt.
                    aktuell.value_pool_entry += code
            if name_teile:
                zusatz = " ".join(name_teile)
                if aktuell.name.endswith("-") and not aktuell.name.endswith(" -"):
                    aktuell.name += zusatz
                else:
                    aktuell.name = (aktuell.name + " " + zusatz).strip()

        for idx, stueck in pid_teile:
            if not stueck:
                continue
            vorher = aktuell.expressions[idx]
            aktuell.expressions[idx] = (vorher + " " + stueck).strip() if vorher else stueck

    for e in entries:
        e.name = " ".join(e.name.split())
        e.expressions = [" ".join(x.split()) for x in e.expressions]
    return entries


def read_document(path: str | Path) -> dict[str, dict]:
    """Liest alle klassischen AHB-Tabellen eines Dokuments."""
    document = docx.Document(str(path))
    ergebnis: dict[str, dict] = {}
    # Je Prüf-ID kann derselbe Tabellenblock im Dokument mehrfach stehen (Dubletten).
    # Deshalb werden die Blöcke getrennt gesammelt und am Ende der umfangreichste je
    # Prüf-ID übernommen, statt die Zeilen zu verdoppeln.
    bloecke: dict[str, list[list[dict]]] = {}
    aktiv: dict[str, list[dict]] = {}
    aktueller_kopf: TableHead | None = None
    section_name = ""
    letzte_segment_id: str | None = None
    letzte_info: StructInfo | None = None

    for table in document.tables:
        # Zeilen und Zellen über das Roh-XML ansprechen: die Grid-Rekonstruktion von
        # python-docx (row.cells) liefert bei Tabellen mit verbundenen Zellen den
        # Inhalt der Nachbarspalte und verfälscht dadurch ganze Segmentblöcke.
        trs = table._tbl.findall(W + "tr")  # noqa: SLF001
        if not trs:
            continue
        kopf_tcs = [_Cell(tc, table) for tc in trs[0].findall(W + "tc")]
        if len(kopf_tcs) < 2:
            continue
        kopfzelle = next(
            (c for c in reversed(kopf_tcs) if "Prüfidentifikator" in c.text),
            kopf_tcs[-1],
        )
        kopf = read_head(kopfzelle)
        start = 0
        if kopf is not None:
            aktueller_kopf = kopf
            section_name = ""
            letzte_segment_id = None
            letzte_info = None
            start = 1
            for pruefi in kopf.pruefis:
                ergebnis.setdefault(
                    pruefi,
                    {
                        "meta": {
                            "pruefidentifikator": pruefi,
                            "description": kopf.beschreibung.get(pruefi, ""),
                            "direction": kopf.kommunikation.get(pruefi, ""),
                            "maus_version": "tabs-reader/1.0",
                        },
                        "lines": [],
                    },
                )
                neuer_block: list[dict] = []
                bloecke.setdefault(pruefi, []).append(neuer_block)
                aktiv[pruefi] = neuer_block
            # Spaltenpositionen an den Inhaltszeilen dieser Tabelle ausrichten
            kalibriere(
                (
                    _Cell(tc, table)
                    for tr in trs[1:40]
                    for tc in tr.findall(W + "tc")[1:2]
                ),
                kopf,
            )
        if aktueller_kopf is None:
            continue

        for tr in trs[start:]:
            zellen = [_Cell(tc, table) for tc in tr.findall(W + "tc")]
            if len(zellen) < 2:
                continue
            struct_zelle = zellen[0]
            body_zelle = zellen[1]
            cond_zelle = zellen[-1] if len(zellen) > 2 else None

            struct_text = struct_zelle.text.strip()
            if not struct_text and not body_zelle.text.strip():
                continue

            # Bei Seitenumbrüchen wiederholt sich der Tabellenkopf mitten in der Tabelle.
            if struct_text.startswith("EDIFACT Struktur"):
                continue

            fortsetzung = False
            info = parse_struct(struct_zelle)
            if info.segment_code is None and info.segment_group_key is None and info.data_element is None:
                if struct_text:
                    section_name = " ".join(struct_text.split())
                    letzte_info = None
                    continue
                # Leere Strukturspalte mit Inhalt: Fortsetzung der vorherigen Zeile
                # (Seitenumbruch innerhalb eines Datenelements).
                if letzte_info is None:
                    continue
                info = letzte_info
                fortsetzung = True
            else:
                letzte_info = info
            if info.segment_id:
                letzte_segment_id = info.segment_id
            elif info.data_element:
                info.segment_id = letzte_segment_id

            eintraege = parse_body(body_zelle, aktueller_kopf)
            bedingung = cond_zelle.text.strip() if cond_zelle is not None else ""

            # Reine Textfortsetzung nach einem Seitenumbruch: die Bezeichnung der
            # zuvor geschriebenen Zeile wird fortgesetzt, es entsteht keine neue Zeile.
            if fortsetzung and eintraege:
                erster = eintraege[0]
                if erster.value_pool_entry is None and not any(erster.expressions) and erster.name:
                    for pruefi in aktueller_kopf.pruefis:
                        zeilen = aktiv.get(pruefi)
                        if not zeilen:
                            continue
                        letzte = zeilen[-1]
                        # nur fortsetzen, wenn die letzte Zeile zum selben Datenelement gehört
                        if (
                            letzte["segment_code"] != info.segment_code
                            or letzte["data_element"] != info.data_element
                            or letzte["segment_group_key"] != info.segment_group_key
                        ):
                            continue
                        vorherige = letzte["name"]
                        trenner = "" if vorherige.endswith("-") else " "
                        letzte["name"] = (vorherige + trenner + erster.name).strip()
                    eintraege = eintraege[1:]

            for eintrag in eintraege:
                for pruefi, expr in zip(aktueller_kopf.pruefis, eintrag.expressions):
                    if not expr:
                        continue
                    ziel = aktiv.get(pruefi)
                    if ziel is None:
                        continue
                    ziel.append(
                        {
                            "ahb_expression": expr,
                            "conditions": bedingung,
                            "data_element": info.data_element,
                            "guid": str(uuid.uuid5(uuid.NAMESPACE_URL, f"{pruefi}|{len(ziel)}|{info.segment_code}|{info.data_element}|{eintrag.value_pool_entry}")),
                            "index": len(ziel) + 1,
                            "name": eintrag.name,
                            "section_name": section_name,
                            "segment_code": info.segment_code,
                            "segment_group_key": info.segment_group_key,
                            "segment_id": info.segment_id,
                            "value_pool_entry": eintrag.value_pool_entry,
                        }
                    )

    for pruefi, kandidaten in bloecke.items():
        beste = max(kandidaten, key=len)
        for i, zeile in enumerate(beste, start=1):
            zeile["index"] = i
        ergebnis[pruefi]["lines"] = beste

    return {p: v for p, v in ergebnis.items() if v["lines"]}


if __name__ == "__main__":
    import sys

    for pfad in sys.argv[1:]:
        res = read_document(pfad)
        print(f"{Path(pfad).name}: {len(res)} Prüf-IDs")
        for pruefi, daten in sorted(res.items())[:10]:
            print(f"   {pruefi}: {len(daten['lines'])} Zeilen  {daten['meta']['description'][:40]!r}")
