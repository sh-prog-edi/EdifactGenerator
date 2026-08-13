#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
lies_aperak_ahb.py — liest die Fehlercodes und Auswertungsregeln der APERAK aus
dem BDEW-AHB APERAK und erzeugt `_engine/daten/aperak-ahb.js`.

Hintergrund (Protokoll Abschnitt 78): Die negative APERAK meldet einen
Verarbeitbarkeitsfehler über SG4 ERC DE9321 (Codes Z10…Z44) und beschreibt den
Fehlerort über Freitextsegmente. Anders als bei der CONTRL gibt es keinen
numerischen Segment-/DE-Zeiger — der AHB definiert dafür aber ein AUSWERTBARES
Format der „Ortsangabe des AHB-Fehlers":

  FTX+Z02, erstes DE4440 : Bezeichnung des Segments, genau wie in der Spalte
                           „Name" der Nachrichtenbeschreibung.
  FTX+Z02, zweites DE4440: das fehlerhafte Segment aus dem Geschäftsvorfall —
                           „beginnt immer mit der Segmentbezeichnung und umfasst
                           alle Zeichen bis ausschließlich dem Segment-Endezeichen"
                           (MIG APERAK zu SG5 FTX+Z02).

Das zweite DE4440 ist damit roher Segmenttext und lässt sich unmittelbar in der
Ursprungsnachricht wiederfinden — die Grundlage für eine echte Positionsanzeige.

Der AHB nennt außerdem die sieben Fehlercodes, bei denen die Ortsangabe Pflicht
ist (Z21, Z29, Z35, Z38, Z39, Z40, Z41), und verknüpft über Bedingungen, welches
FTX bei welchem ERC-Code erwartet wird.

Aufruf:  python3 werkzeuge/lies_aperak_ahb.py [--pfad <AHB_APERAK.docx>]
Ohne --pfad wird der AHB im Arbeitsordner gesucht (EDIGEN_ARBEITSORDNER).
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path

try:
    import docx  # python-docx
except ImportError:  # pragma: no cover
    sys.exit("python-docx fehlt: pip install python-docx")

REPO = Path(__file__).resolve().parents[1]
ARBEITSORDNER = Path(os.environ.get("EDIGEN_ARBEITSORDNER", REPO.parent.parent))
ZIEL = REPO / "_engine" / "daten" / "aperak-ahb.js"

# Zeilenkennung der Strukturtabelle: "SG4\tERC\t9321" bzw. "SG5\tFTX\t00017".
# Die Kennung steht am ZEILENANFANG, gefolgt vom Inhalt der weiteren Spalten.
KENNUNG = re.compile(r"^(SG\d\t)?([A-Z]{3})\t(\d{4,5})(?:\t|$)")
# Ein Codeeintrag der ERC-Liste: "Z29\tErforderliche Angabe für\tX [500]" —
# Code, Bezeichnung (ggf. über Folgezeilen umbrochen), Verwendungsmarke.
# Der erste Code folgt unmittelbar auf die DE-Nummer ("…9321\tZ10\t…"), die
# weiteren stehen nach einem Zeilenumbruch — beide Trenner zulassen.
CODE = re.compile(r"(?:^|\n|\t)(Z\d{2})\t")
BEDINGUNG = re.compile(r"\[(\d+)\]\s*([^\[]+)")


def zeilentext(row) -> str:
    return "\t".join(dict.fromkeys(c.text for c in row.cells))


def finde_ahb(pfad_arg: str | None) -> Path:
    if pfad_arg:
        p = Path(pfad_arg)
        if not p.is_file():
            sys.exit(f"AHB nicht gefunden: {p}")
        return p
    treffer = sorted(ARBEITSORDNER.rglob("AHB_APERAK*.docx"))
    if not treffer:
        sys.exit(
            "AHB_APERAK*.docx nicht im Arbeitsordner gefunden.\n"
            f"  Gesucht unter: {ARBEITSORDNER}\n"
            "  Abhilfe: EDIGEN_ARBEITSORDNER setzen oder --pfad angeben."
        )
    return treffer[-1]


def entzeilen(text: str) -> str:
    """Layout-Umbrüche zu einer Zeile. Bindestriche werden NICHT zusammengezogen —
    im AHB stehen echte Bindestriche mitten in Bezeichnungen ("Segment- bzw.
    Segmentgruppenwiederholbarkeit"), die dabei verstümmelt würden."""
    text = re.sub(r"\s*\n\s*", " ", text)
    return re.sub(r"\s{2,}", " ", text).strip(" \t")


def lies_erc_codes(rows: list[str]) -> list[list[str]]:
    """ERC-Codes DE9321 samt Bezeichnung; Fortsetzung über Seitenumbrüche hinweg."""
    # Der Block beginnt an der Zeile "SG4 ERC 9321" und läuft über die
    # Folgezeilen weiter (dazwischen wiederholte Spaltenköpfe), bis die nächste
    # Segment-/DE-Kennung erscheint.
    start = next((i for i, t in enumerate(rows)
                  if KENNUNG.match(t) and "ERC\t9321" in t), None)
    if start is None:
        return []
    teile = [rows[start]]
    for i in range(start + 1, len(rows)):
        if KENNUNG.match(rows[i]):
            break
        # Nur Fortsetzungszeilen der Codeliste übernehmen. Zwischen den Blöcken
        # stehen Spaltenköpfe (Seitenumbruch) und Strukturzeilen ("Freier Text",
        # "SG4"), die sonst als Bezeichnung des letzten Codes anhängen würden.
        if not re.search(r"Z\d{2}\t", rows[i]):
            continue
        teile.append(rows[i])
    block = "\n".join(teile)

    treffer = list(CODE.finditer(block))
    codes: list[list[str]] = []
    for k, m in enumerate(treffer):
        ende = treffer[k + 1].start() if k + 1 < len(treffer) else len(block)
        roh = block[m.end():ende]
        # Verwendungsmarke ("X", ggf. mit Bedingung) abtrennen — sie steht in
        # einer eigenen Spalte hinter der Bezeichnung.
        marke = ""
        mm = re.search(r"\tX(\s*\[\d+\])?", roh)
        if mm:
            marke = ("X" + (mm.group(1) or "")).strip()
            roh = roh[:mm.start()] + roh[mm.end():]
        # Die Bedingungsspalte ("[500] Hinweis: Für Folgeprozesse.") steht als
        # eigene Zelle MITTEN im Eintrag — sie muss herausgeschnitten werden,
        # nicht abgeschnitten, sonst geht die Fortsetzung der Bezeichnung
        # dahinter verloren (betraf Z14 „Objekt im IT-System nicht gefunden").
        roh = re.sub(r"\[\d+\]\s*Hinweis:[^\[]*?\.", " ", roh, flags=re.S)
        text = entzeilen(roh)
        if text:
            codes.append([m.group(1), text, marke])
    gesehen: dict[str, list[str]] = {}
    for c in codes:
        gesehen.setdefault(c[0], c)
    return [gesehen[k] for k in sorted(gesehen)]


def lies_bedingungen(rows: list[str]) -> dict[str, str]:
    bed: dict[str, str] = {}
    for t in rows:
        for nr, txt in BEDINGUNG.findall(t):
            sauber = entzeilen(txt).rstrip(". ")
            if sauber:
                bed.setdefault(nr, sauber + ".")
    return bed


def lies_ortsangabe_codes(dokument) -> tuple[list[str], dict[str, str]]:
    """Codes mit Pflicht-Ortsangabe — und ihre Bezeichnung aus dem Fließtext.

    Der Fließtext führt dieselben Codes noch einmal als Aufzählung, dort aber
    ohne die Layout-Umbrüche der Tabellenzelle. Er dient deshalb als
    Korrekturquelle für Bezeichnungen, die in der Tabelle mitten im Wort
    umbrechen (z. B. „Segmentgruppenwiederh olbarkeit").
    """
    codes: list[str] = []
    texte: dict[str, str] = {}
    sammeln = False
    for p in dokument.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if "müssen im FTX-Segment" in t or "Ortsangabe des AHB-Fehlers" in t:
            sammeln = True
            continue
        if sammeln:
            m = re.match(r"^(Z\d{2})\s+(.+)$", t)
            if m:
                code = m.group(1)
                if code not in codes:
                    codes.append(code)
                # "… überschritten oder" — das Aufzählungs-"oder" gehört nicht dazu.
                texte.setdefault(code, re.sub(r"\s+oder$", "", m.group(2)).strip())
                continue
            if codes:            # Aufzählung zu Ende
                sammeln = False
    return codes, texte


def schreibe(codes, bedingungen, ortsangabe, quelle) -> None:
    breite = max((len(c[0]) for c in codes), default=4) + 2

    zeilen = [
        "// _engine/daten/aperak-ahb.js",
        "// " + "-" * 66,
        "// Fehlercodes und Auswertungsregeln der APERAK (SG4 ERC DE9321) aus dem",
        "// BDEW-AHB APERAK.",
        "//",
        "// MASCHINELL ERZEUGT durch werkzeuge/lies_aperak_ahb.py — nicht von Hand pflegen.",
        f"// Quelle: {quelle}",
        "//",
        "// Anders als die CONTRL kennt die APERAK keinen numerischen Segment-/DE-Zeiger.",
        "// Der Fehlerort steht als Freitext im FTX+Z02 (Ortsangabe des AHB-Fehlers):",
        "//   erstes DE4440  = Segmentbezeichnung laut Nachrichtenbeschreibung,",
        "//   zweites DE4440 = das fehlerhafte Segment im Rohtext (ohne Endezeichen).",
        "// Das zweite DE4440 ist damit unmittelbar in der Ursprungsnachricht auffindbar.",
        "// " + "-" * 66,
        "var aperakAhb = {",
        '  // [code, Bezeichnung, Verwendungsmarke laut AHB]',
        '  "fehlercodes": [',
    ]
    for i, c in enumerate(codes):
        komma = "," if i + 1 < len(codes) else ""
        felder = f'{json.dumps(c[0], ensure_ascii=False):<{breite}}, ' \
                 f'{json.dumps(c[1], ensure_ascii=False)}, ' \
                 f'{json.dumps(c[2], ensure_ascii=False)}'
        zeilen.append(f"    [{felder}]{komma}")
    zeilen.append("  ],")
    zeilen.append("  // Codes, bei denen der AHB eine Ortsangabe (FTX+Z02) verlangt.")
    zeilen.append('  "mitOrtsangabe": ' + json.dumps(ortsangabe, ensure_ascii=False) + ",")
    zeilen.append('  "bedingungen": {')
    nummern = sorted(bedingungen, key=int)
    for i, nr in enumerate(nummern):
        komma = "," if i + 1 < len(nummern) else ""
        zeilen.append(f'    "{nr}": ' + json.dumps(bedingungen[nr], ensure_ascii=False) + komma)
    zeilen.append("  }")
    zeilen.append("};")
    zeilen.append("")
    zeilen.append("// Klartext zu einem ERC-Fehlercode; leer, wenn der AHB ihn nicht führt.")
    zeilen.append("function aperakFehlertext(code) {")
    zeilen.append("  var liste = (aperakAhb || {}).fehlercodes || [];")
    zeilen.append("  for (var i = 0; i < liste.length; i++)")
    zeilen.append("    if (liste[i][0] === String(code)) return liste[i][1];")
    zeilen.append('  return "";')
    zeilen.append("}")
    zeilen.append("")
    zeilen.append("// Verlangt der AHB zu diesem Code eine Ortsangabe (FTX+Z02)?")
    zeilen.append("function aperakBrauchtOrtsangabe(code) {")
    zeilen.append("  return ((aperakAhb || {}).mitOrtsangabe || []).indexOf(String(code)) >= 0;")
    zeilen.append("}")
    zeilen.append("")
    zeilen.append('if (typeof module !== "undefined")')
    zeilen.append("  module.exports = {")
    zeilen.append("    aperakAhb: aperakAhb,")
    zeilen.append("    aperakFehlertext: aperakFehlertext,")
    zeilen.append("    aperakBrauchtOrtsangabe: aperakBrauchtOrtsangabe,")
    zeilen.append("  };")
    zeilen.append("")
    ZIEL.write_text("\n".join(zeilen), encoding="utf-8")


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--pfad", help="Pfad zur AHB_APERAK-DOCX (sonst Suche im Arbeitsordner)")
    args = ap.parse_args()

    ahb = finde_ahb(args.pfad)
    dokument = docx.Document(str(ahb))
    tabelle = max(dokument.tables, key=lambda t: len(t.rows))
    rows = [zeilentext(r) for r in tabelle.rows]

    codes = lies_erc_codes(rows)
    if not codes:
        sys.exit("Keine ERC-Codes (SG4 ERC DE9321) gefunden — Layout geprüft?")
    bedingungen = lies_bedingungen(rows)
    ortsangabe, fliesstexte = lies_ortsangabe_codes(dokument)
    # Fließtextfassung bevorzugen, wo vorhanden (siehe lies_ortsangabe_codes).
    for c in codes:
        if c[0] in fliesstexte and fliesstexte[c[0]]:
            c[1] = fliesstexte[c[0]]
    schreibe(codes, bedingungen, ortsangabe, ahb.name)

    print(f"Quelle          : {ahb}")
    print(f"ERC-Fehlercodes : {len(codes)}")
    for c in codes:
        print(f"  {c[0]:<5} {c[1][:74]}")
    print(f"mit Ortsangabe  : {', '.join(ortsangabe) or '—'}")
    print(f"Bedingungen     : {len(bedingungen)}")
    print(f"Ziel            : {ZIEL.relative_to(REPO)}")


if __name__ == "__main__":
    main()
